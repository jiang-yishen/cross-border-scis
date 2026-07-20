import paramiko

client = paramiko.SSHClient()
client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
client.connect('47.239.173.51', username='root', password='JWYjys5210@', timeout=30)

# 检查 python-docx 是否可用
stdin, stdout, stderr = client.exec_command('/opt/scis/venv/bin/python -c "import docx; print(\"OK\")"')
out = stdout.read().decode('utf-8').strip()
if out != "OK":
    print("安装 python-docx...")
    client.exec_command('/opt/scis/venv/bin/pip install python-docx -q')

# 创建文档追加脚本
append_script = '''
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('/opt/scis/跨境海外仓供应链智能决策系统_项目汇报.docx')

# 添加新章节：Phase 1-7 优化总结
doc.add_page_break()

# 标题
h1 = doc.add_heading('附件：前端UI优化与性能优化迭代记录（Phase 1-7）', level=1)

# Phase 1
h2 = doc.add_heading('Phase 1：全局样式+导航栏升级', level=2)
p = doc.add_paragraph()
p.add_run('侧边栏宽度：').bold = True
p.add_run('默认180px → 加宽至225-230px，适配大卡片导航按钮')
p = doc.add_paragraph()
p.add_run('导航按钮样式：').bold = True
p.add_run('默认Radio → 大卡片样式（渐变背景+左侧竖条+悬停动效+选中态白色竖条）')
p = doc.add_paragraph()
p.add_run('导航分组：').bold = True
p.add_run('核心运营 / 物流调度 / 系统管理 三大分组，标题居中+两侧装饰线')
p = doc.add_paragraph()
p.add_run('Logo区域：').bold = True
p.add_run('深蓝渐变卡片+品牌名+版本号 SC-Decision Engine v1.0')
p = doc.add_paragraph()
p.add_run('CSS选择器策略：').bold = True
p.add_run('从父子选择器 `.nav-card-wrapper .stButton > button` 改为直接选择器 `[data-testid="stSidebar"] button`，解决Streamlit组件独立渲染问题')

# Phase 2
h2 = doc.add_heading('Phase 2：首页Dashboard+数据可视化升级', level=2)
p = doc.add_paragraph()
p.add_run('KPI卡片：').bold = True
p.add_run('新增左侧4px彩色竖条+图标背景+阴影悬停效果')
p = doc.add_paragraph()
p.add_run('图表容器：').bold = True
p.add_run('统一圆角12px+阴影+悬停放大')
p = doc.add_paragraph()
p.add_run('Plotly主题统一：').bold = True
p.add_run('字体/边距/网格线/图例位置标准化')

# Phase 3-5
h2 = doc.add_heading('Phase 3-5：各页面标题+内容区域升级', level=2)
p = doc.add_paragraph()
p.add_run('统一标题样式：').bold = True
p.add_run('白色卡片+左侧4px竖条+右侧大图标，替代默认st.title()')
p = doc.add_paragraph()
p.add_run('Phase 3：').bold = True
p.add_run('需求预测（预测趋势图+MAPE分布+TOP10排名）+ 补货计划（补货量对比+SKU明细）')
p = doc.add_paragraph()
p.add_run('Phase 4：').bold = True
p.add_run('库存监控（ABC-XYZ散点图+库龄分布+预警表格）+ 调拨建议（优先级柱状图+热力图）+ 采购物流（状态漏斗+在途饼图）')
p = doc.add_paragraph()
p.add_run('Phase 5：').bold = True
p.add_run('使用指南（操作手册下载+导航卡片+系统亮点）+ 运维中心（管理员登录+问题反馈+工单管理+通知发布）')

# Phase 6
h2 = doc.add_heading('Phase 6：登录验证系统（权限控制）', level=2)
p = doc.add_paragraph()
p.add_run('登录页面：').bold = True
p.add_run('左侧蓝白渐变装饰区+右侧登录表单，未登录时CSS隐藏侧边栏')
p = doc.add_paragraph()
p.add_run('预置账户：').bold = True
p.add_run('admin（全权限）/ planner（无运维中心）/ viewer（仅首页+使用指南）')
p = doc.add_paragraph()
p.add_run('权限过滤：').bold = True
p.add_run('基于ROLE_PERMISSIONS字典，在导航渲染时动态过滤可见页面')
p = doc.add_paragraph()
p.add_run('退出登录：').bold = True
p.add_run('导航栏显示用户卡片（头像+用户名+角色）+ "🚪 退出"按钮，一键清除session_state')

# Phase 7
h2 = doc.add_heading('Phase 7：性能测试与优化', level=2)
p = doc.add_paragraph()
p.add_run('数据缓存：').bold = True
p.add_run('st.cache_data(ttl=300) 包装所有数据加载函数（load_sales_daily/load_sku_master/load_replenishment_plan等），缓存时间5分钟，减少重复文件读取')
p = doc.add_paragraph()
p.add_run('KPI计算缓存：').bold = True
p.add_run('st.cache_data(ttl=60) 缓存 compute_kpis()，减少频繁计算')
p = doc.add_paragraph()
p.add_run('并发压测：').bold = True
p.add_run('Locust脚本模拟50用户并发，首页访问权重3/预测页权重2/补货页权重1，持续60秒，输出响应时间报告')

doc.save('/opt/scis/跨境海外仓供应链智能决策系统_项目汇报.docx')
print('OK: docx updated')
'''

sftp = client.open_sftp()
with sftp.file('/tmp/append_docx.py', 'w') as f:
    f.write(append_script.encode('utf-8'))

stdin, stdout, stderr = client.exec_command('/opt/scis/venv/bin/python /tmp/append_docx.py')
out = stdout.read().decode('utf-8')
err = stderr.read().decode('utf-8')
print(out)
if err.strip():
    print("ERR:", err)

client.close()
